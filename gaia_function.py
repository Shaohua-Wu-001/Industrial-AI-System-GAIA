#!/usr/bin/env python3
"""
GAIA Function Tools - Production Grade (Complete for Level 3)
完整支援 GAIA Level 3 所有題目

Version: 2.3.3 Production
Date: 2026-01-06
Security: P0/P1 critical fixes applied with documented limitations
"""

import json
import csv
import base64
import os
import re
import urllib.parse
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
from datetime import datetime, timedelta
from collections import Counter
import ast
import operator
import statistics
import difflib
import ipaddress
import zipfile
import tempfile
import shutil
import math
import socket


# ============================================================
# 配置常數
# ============================================================

MAX_POW_EXP = 1000
MAX_INT_BITS = 20000
MAX_EXPRESSION_LENGTH = 10000
MAX_TEXT_LENGTH = 10000000  # 10MB
MAX_DOWNLOAD_SIZE = 5000000  # 5MB for web_fetch
MAX_ZIP_SIZE = 50000000  # 50MB for zip extraction
MAX_XML_SIZE = 20000000  # 20MB for XML files
MAX_REGEX_TEXT_LENGTH = 200000  # 200KB for regex (reduced for safety)
MAX_PATTERN_LENGTH = 200  # Max regex pattern length
MAX_CSV_EXCEL_SIZE = 50000000  # 50MB for CSV/Excel files
MAX_XML_DEPTH = 200  # Maximum XML nesting depth
MAX_XML_NODES = 100000  # Maximum XML nodes

ALLOWED_SCHEMES = {'http', 'https'}
ALLOWED_PORTS = {None, 80, 443}

# SSRF 防護策略：DNS 失敗時是否放行（預設 fail-closed）
STRICT_SSRF = os.getenv("STRICT_SSRF", "1") not in ("0", "false", "False")

# API 域名白名單
ALLOWED_API_DOMAINS = {
    'api.exchangerate-api.com',
    'nominatim.openstreetmap.org',
}

# Wikipedia 域名白名單
WIKIPEDIA_DOMAIN_SUFFIX = '.wikipedia.org'

# 完整的私有/保留 IP 範圍（包含特殊用途段）
PRIVATE_IP_RANGES = [
    ipaddress.ip_network('0.0.0.0/8'),
    ipaddress.ip_network('10.0.0.0/8'),
    ipaddress.ip_network('100.64.0.0/10'),  # CGNAT
    ipaddress.ip_network('127.0.0.0/8'),
    ipaddress.ip_network('169.254.0.0/16'),
    ipaddress.ip_network('172.16.0.0/12'),
    ipaddress.ip_network('192.0.0.0/24'),   # IETF Protocol Assignments
    ipaddress.ip_network('192.0.2.0/24'),   # TEST-NET-1
    ipaddress.ip_network('192.168.0.0/16'),
    ipaddress.ip_network('198.18.0.0/15'),  # Benchmark
    ipaddress.ip_network('198.51.100.0/24'),  # TEST-NET-2
    ipaddress.ip_network('203.0.113.0/24'),   # TEST-NET-3
    ipaddress.ip_network('224.0.0.0/4'),    # Multicast
    ipaddress.ip_network('240.0.0.0/4'),    # Reserved
    ipaddress.ip_network('255.255.255.255/32'),  # Broadcast
    ipaddress.ip_network('::1/128'),
    ipaddress.ip_network('fc00::/7'),
    ipaddress.ip_network('fe80::/10'),
]


# ============================================================
# 輔助函數
# ============================================================

def _keyify(item: Any) -> str:
    """統一unhashable序列化"""
    try:
        return json.dumps(item, sort_keys=True, ensure_ascii=False)
    except Exception:
        return repr(item)


def _check_int_size(x: Any) -> None:
    """檢查整數大小"""
    if isinstance(x, int) and x.bit_length() > MAX_INT_BITS:
        raise ValueError(f"Integer too large (>{MAX_INT_BITS} bits)")


def _is_within_directory(base_dir: str, target_path: str) -> bool:
    """檢查目標路徑是否在基礎目錄內（使用 relative_to 更穩健）"""
    try:
        base = Path(base_dir).resolve()
        target = Path(target_path).resolve()
        target.relative_to(base)
        return True
    except (ValueError, Exception):
        return False


def _is_zipinfo_symlink(info: zipfile.ZipInfo) -> bool:
    """檢查 ZipInfo 是否為符號連結（防止 symlink 攻擊）"""
    # Unix mode 存在 external_attr 高 16 bits
    mode = (info.external_attr >> 16) & 0o170000
    return mode == 0o120000  # S_IFLNK


def _is_domain_allowed(hostname: str, allowlist: set) -> bool:
    """檢查域名是否在白名單中（支持 subdomain）"""
    hostname_lower = hostname.lower()
    for allowed in allowlist:
        if hostname_lower == allowed or hostname_lower.endswith('.' + allowed):
            return True
    return False


def _is_safe_url(url: str, domain_allowlist: Optional[set] = None) -> Tuple[bool, Optional[str]]:
    """
    SSRF防護：檢查URL安全性
    
    Note: DNS rebinding / TOCTOU not fully mitigated. This is best-effort protection.
    """
    try:
        parsed = urllib.parse.urlparse(url)

        if parsed.scheme not in ALLOWED_SCHEMES:
            return False, f"Scheme '{parsed.scheme}' not allowed"

        if parsed.username or parsed.password:
            return False, "URL userinfo not allowed"

        if parsed.port is not None and parsed.port not in ALLOWED_PORTS:
            return False, f"Port {parsed.port} not allowed (only 80/443)"

        hostname = parsed.hostname
        if not hostname:
            return False, "No hostname"

        # 域名白名單檢查（優先級最高）
        if domain_allowlist and not _is_domain_allowed(hostname, domain_allowlist):
            return False, f"Domain not in allowlist: {hostname}"

        try:
            ip = ipaddress.ip_address(hostname)
            for network in PRIVATE_IP_RANGES:
                if ip in network:
                    return False, f"Private/Reserved IP not allowed: {hostname}"
        except ValueError:
            # hostname 不是 IP，需要 DNS 解析
            try:
                addrs = socket.getaddrinfo(hostname, None, socket.AF_UNSPEC, socket.SOCK_STREAM)
                for addr in addrs:
                    ip_str = addr[4][0]
                    try:
                        ip = ipaddress.ip_address(ip_str)
                        for network in PRIVATE_IP_RANGES:
                            if ip in network:
                                return False, f"Domain resolves to private IP: {ip_str}"
                    except ValueError:
                        pass
            except Exception as dns_err:
                if STRICT_SSRF:
                    return False, f"DNS resolution failed: {str(dns_err)}"

        return True, None
    except Exception as e:
        return False, f"URL parse error: {str(e)}"


def _safe_extract_zip(zip_ref: zipfile.ZipFile, extract_to: str, password: Optional[str] = None):
    """
    安全解壓 ZIP（防止 Zip Slip 和 Symlink 攻擊）
    """
    base = Path(extract_to).resolve()

    for info in zip_ref.infolist():
        # B1 修正：先檢查 symlink（在 is_dir 之前）
        if _is_zipinfo_symlink(info):
            raise ValueError(f"Symlink entry not allowed: {info.filename}")
        
        # 跳過目錄
        if info.is_dir():
            continue

        # 防止絕對路徑、Windows 盤符、.. 穿越
        name = info.filename.replace("\\", "/")
        
        if name.startswith("/"):
            raise ValueError(f"Absolute path in zip not allowed: {info.filename}")
        
        if re.match(r"^[a-zA-Z]:", name):
            raise ValueError(f"Drive letter in zip not allowed: {info.filename}")
        
        dest = (base / name).resolve()
        if not _is_within_directory(str(base), str(dest)):
            raise ValueError(f"Zip Slip detected: {info.filename}")

    # 檢查完畢後再 extract
    if password:
        zip_ref.extractall(extract_to, pwd=password.encode())
    else:
        zip_ref.extractall(extract_to)


def _create_safe_session() -> Any:
    """創建安全的 requests session（禁用環境代理）"""
    try:
        import requests
        session = requests.Session()
        session.trust_env = False
        return session
    except ImportError:
        return None


def _build_line_index(text: str) -> List[int]:
    """預建行索引"""
    line_starts = [0]
    for i, char in enumerate(text):
        if char == '\n':
            line_starts.append(i + 1)
    return line_starts


def _find_line_number(position: int, line_starts: List[int]) -> int:
    """二分查找行號"""
    import bisect
    idx = bisect.bisect_right(line_starts, position)
    return idx


def _rankdata(values: List[float]) -> List[float]:
    """Spearman用：處理ties的平均名次"""
    n = len(values)
    sorted_idx = sorted(range(n), key=lambda i: values[i])
    ranks = [0.0] * n
    i = 0
    while i < n:
        j = i
        while j + 1 < n and values[sorted_idx[j + 1]] == values[sorted_idx[i]]:
            j += 1
        avg_rank = (i + j) / 2.0 + 1.0
        for k in range(i, j + 1):
            ranks[sorted_idx[k]] = avg_rank
        i = j + 1
    return ranks


# ============================================================
# Priority 1 Tools (原始17個)
# ============================================================

def read_pdf(file_path: str, page_numbers: Optional[List[int]] = None) -> Dict[str, Any]:
    """Read PDF file. Page numbers are 1-indexed."""
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)

            metadata = {}
            if pdf_reader.metadata:
                metadata = {
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                    "creator": pdf_reader.metadata.get("/Creator", "")
                }

            content = ""
            if page_numbers:
                for page_num in page_numbers:
                    if 1 <= page_num <= total_pages:
                        page = pdf_reader.pages[page_num - 1]
                        txt = page.extract_text() or ""
                        content += txt + "\n\n"
            else:
                for page in pdf_reader.pages:
                    txt = page.extract_text() or ""
                    content += txt + "\n\n"

        return {
            "success": True,
            "content": content.strip(),
            "pages": total_pages,
            "metadata": metadata,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "content": None,
            "pages": 0,
            "metadata": {},
            "error": str(e)
        }


def web_search(query: str, num_results: int = 5) -> Dict[str, Any]:
    """Web search with Serper API or simulation."""
    api_key = os.environ.get("SERPER_API_KEY")

    if not api_key:
        results = [
            {
                "title": f"Search result {i+1} - {query}",
                "url": f"https://example.com/result{i+1}",
                "snippet": f"Summary for '{query}' result {i+1}"
            }
            for i in range(min(num_results, 5))
        ]
        return {
            "success": True,
            "results": results,
            "query": query,
            "count": len(results),
            "provider": "Simulation",
            "is_simulated": True,
            "error": None
        }

    try:
        import requests
        url = "https://google.serper.dev/search"
        payload = json.dumps({"q": query, "num": num_results})
        headers = {
            'X-API-KEY': api_key,
            'Content-Type': 'application/json'
        }

        session = _create_safe_session()
        if session:
            response = session.post(url, headers=headers, data=payload, timeout=10)
        else:
            response = requests.post(url, headers=headers, data=payload, timeout=10, 
                                    proxies={"http": None, "https": None})
        response.raise_for_status()

        data = response.json()
        results = []

        for item in data.get("organic", [])[:num_results]:
            results.append({
                "title": item.get("title", ""),
                "url": item.get("link", ""),
                "snippet": item.get("snippet", "")
            })

        return {
            "success": True,
            "results": results,
            "query": query,
            "count": len(results),
            "provider": "Serper (Google)",
            "is_simulated": False,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "results": [],
            "query": query,
            "count": 0,
            "provider": "Serper",
            "is_simulated": False,
            "error": str(e)
        }


def read_csv(file_path: str, encoding: str = "utf-8") -> Dict[str, Any]:
    """Read CSV file with size limit."""
    try:
        # D1: 檢查文件大小
        file_size = os.path.getsize(file_path)
        if file_size > MAX_CSV_EXCEL_SIZE:
            return {
                "success": False,
                "data": None,
                "columns": [],
                "rows": 0,
                "error": f"CSV file too large (>{MAX_CSV_EXCEL_SIZE} bytes)"
            }
        
        import pandas as pd
        df = pd.read_csv(file_path, encoding=encoding)
        return {
            "success": True,
            "data": df.to_dict('records'),
            "columns": list(df.columns),
            "rows": len(df),
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "data": None,
            "columns": [],
            "rows": 0,
            "error": str(e)
        }


def read_image(file_path: str) -> Dict[str, Any]:
    """Read image file."""
    try:
        from PIL import Image
        img = Image.open(file_path)
        return {
            "success": True,
            "format": img.format,
            "mode": img.mode,
            "width": img.width,
            "height": img.height,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "format": None,
            "mode": None,
            "width": 0,
            "height": 0,
            "error": str(e)
        }


def calculate(expression: str) -> Dict[str, Any]:
    """Safe math expression evaluator with pow estimation and finite check."""
    try:
        if len(expression) > MAX_EXPRESSION_LENGTH:
            raise ValueError(f"Expression too long (max {MAX_EXPRESSION_LENGTH})")

        allowed_operators = {
            ast.Add: operator.add,
            ast.Sub: operator.sub,
            ast.Mult: operator.mul,
            ast.Div: operator.truediv,
            ast.Pow: operator.pow,
            ast.USub: operator.neg,
        }

        def eval_node(node):
            if isinstance(node, ast.Constant):
                if not isinstance(node.value, (int, float)):
                    raise ValueError("Only numeric constants allowed")
                _check_int_size(node.value)
                return node.value

            elif isinstance(node, ast.Num):
                _check_int_size(node.n)
                return node.n

            elif isinstance(node, ast.BinOp):
                left = eval_node(node.left)
                right = eval_node(node.right)
                op_type = type(node.op)

                if op_type not in allowed_operators:
                    raise ValueError(f"Unsupported operator: {op_type.__name__}")

                if op_type is ast.Pow:
                    if not isinstance(right, (int, float)):
                        raise ValueError("Exponent must be numeric")
                    if abs(right) > MAX_POW_EXP:
                        raise ValueError(f"Exponent too large (max {MAX_POW_EXP})")

                    if isinstance(left, int) and isinstance(right, int) and right >= 0:
                        est_bits = left.bit_length() * right
                        if est_bits > MAX_INT_BITS:
                            raise ValueError(f"Power result too large (est {est_bits} bits)")

                result = allowed_operators[op_type](left, right)
                
                if isinstance(result, float) and not math.isfinite(result):
                    raise ValueError("Non-finite result (inf/nan)")
                
                _check_int_size(result)
                return result

            elif isinstance(node, ast.UnaryOp):
                operand = eval_node(node.operand)
                op_type = type(node.op)

                if op_type not in allowed_operators:
                    raise ValueError(f"Unsupported unary operator: {op_type.__name__}")

                result = allowed_operators[op_type](operand)
                _check_int_size(result)
                return result

            else:
                raise ValueError("Unsupported operation")

        tree = ast.parse(expression, mode='eval')
        result = eval_node(tree.body)

        return {
            "success": True,
            "result": result,
            "expression": expression,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "result": None,
            "expression": expression,
            "error": str(e)
        }


def read_excel(file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
    """Read Excel file with size limit."""
    try:
        # D1: 檢查文件大小
        file_size = os.path.getsize(file_path)
        if file_size > MAX_CSV_EXCEL_SIZE:
            return {
                "success": False,
                "data": None,
                "columns": [],
                "rows": 0,
                "error": f"Excel file too large (>{MAX_CSV_EXCEL_SIZE} bytes)"
            }
        
        import pandas as pd
        if sheet_name:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
        else:
            df = pd.read_excel(file_path)

        return {
            "success": True,
            "data": df.to_dict('records'),
            "columns": list(df.columns),
            "rows": len(df),
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "data": None,
            "columns": [],
            "rows": 0,
            "error": str(e)
        }


def image_to_text(file_path: str, lang: str = "eng") -> Dict[str, Any]:
    """OCR text extraction."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang=lang)
        return {
            "success": True,
            "text": text.strip(),
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "text": None,
            "error": str(e)
        }


def web_fetch(url: str, timeout: int = 10) -> Dict[str, Any]:
    """Fetch webpage with SSRF protection and proper resource management."""
    try:
        ok, msg = _is_safe_url(url)
        if not ok:
            return {
                "success": False,
                "content": None,
                "status_code": None,
                "error": f"URL blocked: {msg}"
            }

        import requests
        from bs4 import BeautifulSoup

        session = _create_safe_session() or requests.Session()
        current_url = url
        max_redirects = 5

        for _ in range(max_redirects + 1):
            # C1: 使用 context manager 確保連線釋放
            with session.get(
                current_url,
                timeout=(3, timeout),
                allow_redirects=False,
                headers={
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.9",
                "Accept-Encoding": "gzip, deflate, br",
                "DNT": "1",
                "Connection": "keep-alive",
                "Upgrade-Insecure-Requests": "1",
                "Sec-Fetch-Dest": "document",
                "Sec-Fetch-Mode": "navigate",
                "Sec-Fetch-Site": "none",
                "Sec-Fetch-User": "?1",
                "Cache-Control": "max-age=0"
            },
                stream=True
            ) as resp:
                
                if 300 <= resp.status_code < 400 and resp.headers.get("Location"):
                    next_url = urllib.parse.urljoin(current_url, resp.headers["Location"])
                    ok, msg = _is_safe_url(next_url)
                    if not ok:
                        return {
                            "success": False,
                            "content": None,
                            "status_code": resp.status_code,
                            "error": f"Redirect blocked: {msg}"
                        }
                    current_url = next_url
                    continue

                resp.raise_for_status()

                # D3: 更保守的 Content-Type 檢查（空值視為不支援）
                content_type = resp.headers.get("Content-Type", "").lower()
                if not content_type:
                    return {
                        "success": False,
                        "content": None,
                        "status_code": resp.status_code,
                        "error": "Missing Content-Type header"
                    }
                
                if not any(ct in content_type for ct in ["text/", "application/json", "application/xml", "application/xhtml"]):
                    return {
                        "success": False,
                        "content": None,
                        "status_code": resp.status_code,
                        "error": f"Unsupported content type: {content_type}"
                    }

                buf = bytearray()
                for chunk in resp.iter_content(chunk_size=8192):
                    if not chunk:
                        continue
                    buf.extend(chunk)
                    if len(buf) > MAX_DOWNLOAD_SIZE:
                        return {
                            "success": False,
                            "content": None,
                            "status_code": resp.status_code,
                            "error": f"Content too large (>{MAX_DOWNLOAD_SIZE} bytes)"
                        }

                text_content = bytes(buf).decode(resp.encoding or 'utf-8', errors='ignore')
                
                try:
                    soup = BeautifulSoup(text_content, 'html.parser')
                    text = soup.get_text(separator='\n', strip=True)
                except Exception:
                    text = text_content

                return {
                    "success": True,
                    "content": text,
                    "status_code": resp.status_code,
                    "error": None
                }

        return {
            "success": False,
            "content": None,
            "status_code": None,
            "error": "Too many redirects"
        }

    except Exception as e:
        return {
            "success": False,
            "content": None,
            "status_code": None,
            "error": str(e)
        }


def filter_data(data: List[Dict], conditions: Dict[str, Any]) -> Dict[str, Any]:
    """Filter data by conditions."""
    try:
        original_count = len(data)
        filtered = []
        for item in data:
            match = True
            for key, value in conditions.items():
                if key not in item or item[key] != value:
                    match = False
                    break
            if match:
                filtered.append(item)

        return {
            "success": True,
            "filtered_data": filtered,
            "count": len(filtered),
            "original_count": original_count,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "filtered_data": [],
            "count": 0,
            "original_count": 0,
            "error": str(e)
        }


def read_json(file_path: str, encoding: str = "utf-8") -> Dict[str, Any]:
    """Read JSON file (supports .json and .jsonld)."""
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            data = json.load(f)

        data_type = "dict" if isinstance(data, dict) else "array" if isinstance(data, list) else "other"

        return {
            "success": True,
            "data": data,
            "type": data_type,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "data": None,
            "type": None,
            "error": str(e)
        }


def read_text_file(file_path: str, encoding: str = "utf-8") -> Dict[str, Any]:
    """Read text file with size limit."""
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_TEXT_LENGTH:
            return {
                "success": False,
                "content": None,
                "lines": 0,
                "characters": 0,
                "error": f"Text file too large (>{MAX_TEXT_LENGTH} bytes)"
            }
        
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()

        return {
            "success": True,
            "content": content,
            "lines": len(content.split('\n')),
            "characters": len(content),
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "content": None,
            "lines": 0,
            "characters": 0,
            "error": str(e)
        }


def read_docx(file_path: str) -> Dict[str, Any]:
    """Read Word document."""
    try:
        import docx
        doc = docx.Document(file_path)
        content = '\n\n'.join([para.text for para in doc.paragraphs])
        return {
            "success": True,
            "content": content,
            "paragraphs": len(doc.paragraphs),
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "content": None,
            "paragraphs": 0,
            "error": str(e)
        }


def aggregate_data(data: List[Dict], group_by: str,
                   aggregate_field: str, operation: str = "sum") -> Dict[str, Any]:
    """Aggregate data."""
    try:
        import pandas as pd
        df = pd.DataFrame(data)

        if operation == "sum":
            result = df.groupby(group_by)[aggregate_field].sum()
        elif operation == "mean":
            result = df.groupby(group_by)[aggregate_field].mean()
        elif operation == "count":
            result = df.groupby(group_by)[aggregate_field].count()
        else:
            raise ValueError(f"Unsupported operation: {operation}")

        return {
            "success": True,
            "aggregated_data": result.to_dict(),
            "operation": operation,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "aggregated_data": None,
            "operation": None,
            "error": str(e)
        }


def sort_data(data: List[Dict], sort_by: str,
              reverse: bool = False) -> Dict[str, Any]:
    """Sort data."""
    try:
        sorted_data = sorted(
            data,
            key=lambda x: x.get(sort_by, 0),
            reverse=reverse
        )
        order = "descending" if reverse else "ascending"

        return {
            "success": True,
            "sorted_data": sorted_data,
            "count": len(sorted_data),
            "sort_by": sort_by,
            "order": order,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "sorted_data": [],
            "count": 0,
            "sort_by": None,
            "order": None,
            "error": str(e)
        }


def analyze_image(file_path: str) -> Dict[str, Any]:
    """Analyze image properties."""
    try:
        from PIL import Image
        import numpy as np

        img = Image.open(file_path)
        img_array = np.array(img)

        return {
            "success": True,
            "format": img.format,
            "mode": img.mode,
            "size": img.size,
            "mean_color": img_array.mean(axis=(0,1)).tolist() if len(img_array.shape) == 3 else None,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "format": None,
            "mode": None,
            "size": None,
            "mean_color": None,
            "error": str(e)
        }


def date_calculator(start_date: str, days_to_add: int = 0,
                    end_date: Optional[str] = None,
                    date_format: str = "%Y-%m-%d") -> Dict[str, Any]:
    """Date calculator."""
    try:
        start = datetime.strptime(start_date, date_format)

        result_dict = {
            "success": True,
            "start_date": start_date,
            "days_added": None,
            "result_date": None,
            "end_date": None,
            "days_diff": None,
            "error": None
        }

        if days_to_add != 0:
            result_date = start + timedelta(days=days_to_add)
            result_dict["days_added"] = days_to_add
            result_dict["result_date"] = result_date.strftime(date_format)

        if end_date:
            end = datetime.strptime(end_date, date_format)
            result_dict["end_date"] = end_date
            result_dict["days_diff"] = (end - start).days

        return result_dict

    except Exception as e:
        return {
            "success": False,
            "start_date": start_date,
            "days_added": None,
            "result_date": None,
            "end_date": None,
            "days_diff": None,
            "error": str(e)
        }


def unit_converter(value: float, from_unit: str,
                    to_unit: str, unit_type: str = "length") -> Dict[str, Any]:
    """Unit converter."""
    try:
        conversions = {
            "length": {
                "m": 1.0,
                "km": 0.001,
                "cm": 100.0,
                "mile": 0.000621371,
            },
            "weight": {
                "kg": 1.0,
                "g": 1000.0,
                "lb": 2.20462,
            },
            "volume": {
                "L": 1.0,
                "mL": 1000.0,
                "gallon": 3.78541,
            }
        }

        if unit_type == "temperature":
            if from_unit == "C" and to_unit == "F":
                result = (value * 9/5) + 32
            elif from_unit == "F" and to_unit == "C":
                result = (value - 32) * 5/9
            elif from_unit == "C" and to_unit == "K":
                result = value + 273.15
            elif from_unit == "K" and to_unit == "C":
                result = value - 273.15
            else:
                raise ValueError("Unsupported temperature conversion")
        else:
            if unit_type not in conversions:
                raise ValueError(f"Unknown category: {unit_type}")

            if from_unit not in conversions[unit_type] or to_unit not in conversions[unit_type]:
                raise ValueError(f"Unknown unit")

            base_value = value / conversions[unit_type][from_unit]
            result = base_value * conversions[unit_type][to_unit]

        return {
            "success": True,
            "value": value,
            "from_unit": from_unit,
            "to_unit": to_unit,
            "result": result,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "value": value,
            "from_unit": from_unit,
            "to_unit": to_unit,
            "result": None,
            "error": str(e)
        }


# ============================================================
# 新增工具 (19個)
# ============================================================

def join_data(
    data1: List[Dict[str, Any]],
    data2: List[Dict[str, Any]],
    join_key: str,
    join_type: str = "inner"
) -> Dict[str, Any]:
    """Join data with duplicate key detection."""
    try:
        result = []

        data2_keys = []
        for item in data2:
            if join_key in item and item[join_key] is not None:
                data2_keys.append(item[join_key])

        key_counts = Counter(data2_keys)
        duplicates = [k for k, count in key_counts.items() if count > 1]
        if duplicates:
            return {
                "success": False,
                "data": None,
                "count": 0,
                "join_type": join_type,
                "error": f"Duplicate keys in data2: {duplicates[:5]}"
            }

        data2_dict = {
            item[join_key]: item
            for item in data2
            if join_key in item and item[join_key] is not None
        }

        warnings = []
        if join_type == "right":
            data1_keys = [item[join_key] for item in data1 if join_key in item and item[join_key] is not None]
            data1_key_counts = Counter(data1_keys)
            data1_duplicates = [k for k, count in data1_key_counts.items() if count > 1]
            if data1_duplicates:
                warnings.append(f"Duplicate keys in data1 (first occurrence used): {data1_duplicates[:5]}")

        if join_type == "inner":
            for item1 in data1:
                key = item1.get(join_key)
                if key is not None and key in data2_dict:
                    merged = {**item1, **data2_dict[key]}
                    result.append(merged)

        elif join_type == "left":
            for item1 in data1:
                key = item1.get(join_key)
                if key is not None and key in data2_dict:
                    merged = {**item1, **data2_dict[key]}
                else:
                    merged = item1
                result.append(merged)

        elif join_type == "right":
            data1_dict = {}
            for item in data1:
                if join_key in item and item[join_key] is not None:
                    k = item[join_key]
                    if k not in data1_dict:
                        data1_dict[k] = item

            for item2 in data2:
                key = item2.get(join_key)
                if key is not None and key in data1_dict:
                    merged = {**data1_dict[key], **item2}
                else:
                    merged = item2
                result.append(merged)

        elif join_type == "outer":
            processed_keys = set()
            for item1 in data1:
                key = item1.get(join_key)
                if key is not None and key in data2_dict:
                    merged = {**item1, **data2_dict[key]}
                    result.append(merged)
                    processed_keys.add(key)
                else:
                    result.append(item1)

            for item2 in data2:
                key = item2.get(join_key)
                if key is not None and key not in processed_keys:
                    result.append(item2)

        return_dict = {
            "success": True,
            "data": result,
            "count": len(result),
            "join_type": join_type,
            "error": None
        }
        
        if warnings:
            return_dict["warnings"] = warnings

        return return_dict

    except Exception as e:
        return {
            "success": False,
            "data": None,
            "count": 0,
            "join_type": join_type,
            "error": str(e)
        }


def deduplicate_data(
    data: List[Dict[str, Any]],
    key_fields: Optional[List[str]] = None
) -> Dict[str, Any]:
    """Deduplicate data."""
    try:
        if not key_fields:
            seen = set()
            result = []
            for item in data:
                item_str = _keyify(item)
                if item_str not in seen:
                    seen.add(item_str)
                    result.append(item)
        else:
            seen = set()
            result = []
            for item in data:
                key = tuple(item.get(field) for field in key_fields)
                if key not in seen:
                    seen.add(key)
                    result.append(item)

        original_count = len(data)
        deduplicated_count = len(result)
        removed_count = original_count - deduplicated_count

        return {
            "success": True,
            "data": result,
            "original_count": original_count,
            "deduplicated_count": deduplicated_count,
            "removed_count": removed_count,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "data": None,
            "original_count": 0,
            "deduplicated_count": 0,
            "removed_count": 0,
            "error": str(e)
        }


def pivot_table(
    data: List[Dict[str, Any]],
    index: str,
    values: str,
    aggfunc: str = "sum"
) -> Dict[str, Any]:
    """Create pivot table."""
    try:
        grouped = {}
        for item in data:
            key = item.get(index)
            if key is not None:
                if key not in grouped:
                    grouped[key] = []
                value = item.get(values)
                if value is not None:
                    grouped[key].append(float(value))

        result = {}
        for key, vals in grouped.items():
            if aggfunc == "sum":
                result[key] = sum(vals)
            elif aggfunc == "mean":
                result[key] = sum(vals) / len(vals)
            elif aggfunc == "count":
                result[key] = len(vals)
            elif aggfunc == "min":
                result[key] = min(vals)
            elif aggfunc == "max":
                result[key] = max(vals)
            else:
                raise ValueError(f"Unsupported aggfunc: {aggfunc}")

        return {
            "success": True,
            "pivot_data": result,
            "index": index,
            "values": values,
            "aggfunc": aggfunc,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "pivot_data": None,
            "index": index,
            "values": values,
            "aggfunc": aggfunc,
            "error": str(e)
        }


def fill_missing(
    data: List[Dict[str, Any]],
    columns: List[str],
    method: str = "forward_fill"
) -> Dict[str, Any]:
    """Fill missing values."""
    try:
        result = []

        if method == "forward_fill":
            last_values = {}
            for item in data:
                new_item = item.copy()
                for col in columns:
                    if col in item and item[col] is not None:
                        last_values[col] = item[col]
                    elif col in last_values:
                        new_item[col] = last_values[col]
                result.append(new_item)

        elif method == "backward_fill":
            data_reversed = list(reversed(data))
            last_values = {}
            temp_result = []
            for item in data_reversed:
                new_item = item.copy()
                for col in columns:
                    if col in item and item[col] is not None:
                        last_values[col] = item[col]
                    elif col in last_values:
                        new_item[col] = last_values[col]
                temp_result.append(new_item)
            result = list(reversed(temp_result))

        elif method in ["mean", "median"]:
            col_values = {col: [] for col in columns}
            for item in data:
                for col in columns:
                    if col in item and item[col] is not None:
                        try:
                            col_values[col].append(float(item[col]))
                        except:
                            pass

            fill_values = {}
            for col, values in col_values.items():
                if values:
                    if method == "mean":
                        fill_values[col] = sum(values) / len(values)
                    else:
                        fill_values[col] = statistics.median(values)

            for item in data:
                new_item = item.copy()
                for col in columns:
                    if col not in item or item[col] is None:
                        if col in fill_values:
                            new_item[col] = fill_values[col]
                result.append(new_item)

        return {
            "success": True,
            "data": result,
            "method": method,
            "columns": columns,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "data": None,
            "method": method,
            "columns": columns,
            "error": str(e)
        }


def sample_data(
    data: List[Dict[str, Any]],
    n: int,
    random_seed: Optional[int] = None
) -> Dict[str, Any]:
    """Random sample."""
    try:
        import random

        if random_seed is not None:
            random.seed(random_seed)

        sample_size = min(n, len(data))
        sampled = random.sample(data, sample_size)

        return {
            "success": True,
            "data": sampled,
            "requested_size": n,
            "actual_size": sample_size,
            "original_size": len(data),
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "data": None,
            "requested_size": n,
            "actual_size": 0,
            "original_size": len(data),
            "error": str(e)
        }


def regex_search(
    text: str,
    pattern: str,
    return_all: bool = True
) -> Dict[str, Any]:
    """
    Regex search with ReDoS protection.
    
    Note: ReDoS risk reduced via pattern/text length limits; timeout not enforced.
    """
    try:
        if len(pattern) > MAX_PATTERN_LENGTH:
            return {
                "success": False,
                "matches": [],
                "count": 0,
                "pattern": pattern,
                "error": f"Pattern too long (max {MAX_PATTERN_LENGTH})"
            }
        
        if len(text) > MAX_REGEX_TEXT_LENGTH:
            return {
                "success": False,
                "matches": [],
                "count": 0,
                "pattern": pattern,
                "error": f"Text too large for regex (max {MAX_REGEX_TEXT_LENGTH})"
            }

        if return_all:
            matches = re.findall(pattern, text)
        else:
            match = re.search(pattern, text)
            matches = [match.group()] if match else []

        return {
            "success": True,
            "matches": matches,
            "count": len(matches),
            "pattern": pattern,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "matches": [],
            "count": 0,
            "pattern": pattern,
            "error": str(e)
        }


def string_transform(
    text: str,
    operation: str
) -> Dict[str, Any]:
    """String transformation."""
    try:
        if operation == "upper":
            result = text.upper()
        elif operation == "lower":
            result = text.lower()
        elif operation == "title":
            result = text.title()
        elif operation == "reverse":
            result = text[::-1]
        elif operation == "strip":
            result = text.strip()
        elif operation == "capitalize":
            result = text.capitalize()
        else:
            return {
                "success": False,
                "result": None,
                "operation": operation,
                "error": f"Unknown operation: {operation}"
            }

        return {
            "success": True,
            "result": result,
            "operation": operation,
            "original_length": len(text),
            "result_length": len(result),
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "result": None,
            "operation": operation,
            "error": str(e)
        }


def encode_decode(
    text: str,
    operation: str
) -> Dict[str, Any]:
    """Encode/decode text with error handling."""
    try:
        if operation == "base64_encode":
            result = base64.b64encode(text.encode()).decode()
        elif operation == "base64_decode":
            try:
                decoded_bytes = base64.b64decode(text.encode())
                result = decoded_bytes.decode('utf-8')
            except UnicodeDecodeError:
                result = decoded_bytes.hex()
        elif operation == "url_encode":
            result = urllib.parse.quote(text)
        elif operation == "url_decode":
            result = urllib.parse.unquote(text)
        else:
            return {
                "success": False,
                "result": None,
                "operation": operation,
                "error": f"Unknown operation: {operation}"
            }

        return {
            "success": True,
            "result": result,
            "operation": operation,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "result": None,
            "operation": operation,
            "error": str(e)
        }


def split_join_text(
    text: str,
    operation: str,
    separator: str = ","
) -> Dict[str, Any]:
    """Split/join text."""
    try:
        if operation == "split":
            result = text.split(separator)
        elif operation == "join":
            try:
                parts = json.loads(text)
                result = separator.join(str(p) for p in parts)
            except:
                parts = text.split('\n')
                result = separator.join(parts)
        else:
            return {
                "success": False,
                "result": None,
                "operation": operation,
                "error": f"Unknown operation: {operation}"
            }

        return {
            "success": True,
            "result": result,
            "operation": operation,
            "separator": separator,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "result": None,
            "operation": operation,
            "separator": separator,
            "error": str(e)
        }


def statistical_analysis(
    data: List[float],
    metrics: List[str]
) -> Dict[str, Any]:
    """Statistical analysis."""
    try:
        results = {}

        if "mean" in metrics:
            results["mean"] = statistics.mean(data)
        if "median" in metrics:
            results["median"] = statistics.median(data)
        if "mode" in metrics:
            try:
                results["mode"] = statistics.mode(data)
            except statistics.StatisticsError:
                results["mode"] = None
        if "std" in metrics or "stdev" in metrics:
            results["std"] = statistics.stdev(data) if len(data) > 1 else 0
        if "variance" in metrics:
            results["variance"] = statistics.variance(data) if len(data) > 1 else 0
        if "min" in metrics:
            results["min"] = min(data)
        if "max" in metrics:
            results["max"] = max(data)
        if "sum" in metrics:
            results["sum"] = sum(data)
        if "count" in metrics:
            results["count"] = len(data)
        if "range" in metrics:
            results["range"] = max(data) - min(data)

        for metric in metrics:
            if metric.startswith("percentile_"):
                try:
                    p = int(metric.split("_")[1])
                    sorted_data = sorted(data)
                    k = (len(sorted_data) - 1) * p / 100
                    f = int(k)
                    c = f + 1 if f + 1 < len(sorted_data) else f
                    results[metric] = sorted_data[f] + (k - f) * (sorted_data[c] - sorted_data[f])
                except:
                    pass

        return {
            "success": True,
            "statistics": results,
            "data_size": len(data),
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "statistics": None,
            "data_size": len(data) if data else 0,
            "error": str(e)
        }


def correlation_analysis(
    data: List[Dict[str, Any]],
    x_column: str,
    y_column: str,
    method: str = "pearson"
) -> Dict[str, Any]:
    """Correlation analysis with Spearman support."""
    try:
        x_values = []
        y_values = []

        for item in data:
            if x_column in item and y_column in item:
                try:
                    x_values.append(float(item[x_column]))
                    y_values.append(float(item[y_column]))
                except:
                    pass

        if len(x_values) < 2:
            return {
                "success": False,
                "correlation": None,
                "method": method,
                "sample_size": 0,
                "error": "Not enough valid data points"
            }

        method_l = (method or "").lower()
        if method_l not in ("pearson", "spearman"):
            return {
                "success": False,
                "correlation": None,
                "method": method,
                "sample_size": len(x_values),
                "error": f"Unsupported method: {method}"
            }

        if method_l == "spearman":
            x_use = _rankdata(x_values)
            y_use = _rankdata(y_values)
        else:
            x_use = x_values
            y_use = y_values

        n = len(x_use)
        sum_x = sum(x_use)
        sum_y = sum(y_use)
        sum_xy = sum(x * y for x, y in zip(x_use, y_use))
        sum_x2 = sum(x ** 2 for x in x_use)
        sum_y2 = sum(y ** 2 for y in y_use)

        numerator = n * sum_xy - sum_x * sum_y
        denominator = ((n * sum_x2 - sum_x ** 2) * (n * sum_y2 - sum_y ** 2)) ** 0.5

        correlation = numerator / denominator if denominator != 0 else 0.0

        return {
            "success": True,
            "correlation": correlation,
            "method": method_l,
            "sample_size": n,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "correlation": None,
            "method": method,
            "sample_size": 0,
            "error": str(e)
        }


def moving_average(
    data: List[float],
    window_size: int
) -> Dict[str, Any]:
    """Moving average."""
    try:
        if window_size <= 0 or window_size > len(data):
            return {
                "success": False,
                "result": None,
                "window_size": window_size,
                "original_size": len(data),
                "result_size": 0,
                "error": "Invalid window size"
            }

        result = []
        for i in range(len(data) - window_size + 1):
            window = data[i:i + window_size]
            avg = sum(window) / window_size
            result.append(avg)

        return {
            "success": True,
            "result": result,
            "window_size": window_size,
            "original_size": len(data),
            "result_size": len(result),
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "result": None,
            "window_size": window_size,
            "original_size": len(data) if data else 0,
            "result_size": 0,
            "error": str(e)
        }


def currency_converter(
    amount: float,
    from_currency: str,
    to_currency: str
) -> Dict[str, Any]:
    """Currency converter with SSRF protection and input validation."""
    try:
        if not re.fullmatch(r"[A-Z]{3}", from_currency.upper()):
            return {
                "success": False,
                "result": None,
                "rate": None,
                "amount": amount,
                "from_currency": from_currency,
                "to_currency": to_currency,
                "error": "Invalid from_currency code (must be 3 uppercase letters)"
            }
        
        if not re.fullmatch(r"[A-Z]{3}", to_currency.upper()):
            return {
                "success": False,
                "result": None,
                "rate": None,
                "amount": amount,
                "from_currency": from_currency,
                "to_currency": to_currency,
                "error": "Invalid to_currency code (must be 3 uppercase letters)"
            }

        import requests

        from_currency = from_currency.upper()
        to_currency = to_currency.upper()
        url = f"https://api.exchangerate-api.com/v4/latest/{from_currency}"

        ok, msg = _is_safe_url(url, ALLOWED_API_DOMAINS)
        if not ok:
            return {
                "success": False,
                "result": None,
                "rate": None,
                "amount": amount,
                "from_currency": from_currency,
                "to_currency": to_currency,
                "error": f"URL blocked: {msg}"
            }

        session = _create_safe_session() or requests.Session()
        response = session.get(url, timeout=10)
        response.raise_for_status()

        data = response.json()

        if to_currency not in data['rates']:
            return {
                "success": False,
                "result": None,
                "rate": None,
                "amount": amount,
                "from_currency": from_currency,
                "to_currency": to_currency,
                "error": f"Currency {to_currency} not found"
            }

        rate = data['rates'][to_currency]
        result = amount * rate

        return {
            "success": True,
            "result": result,
            "amount": amount,
            "from_currency": from_currency,
            "to_currency": to_currency,
            "rate": rate,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "result": None,
            "rate": None,
            "amount": amount,
            "from_currency": from_currency,
            "to_currency": to_currency,
            "error": str(e)
        }


def wikipedia_search(
    query: str,
    language: str = "en",
    sentences: int = 3
) -> Dict[str, Any]:
    """Wikipedia search with language validation and domain allowlist."""
    try:
        if not re.fullmatch(r"[a-z]{2,12}(-[a-z0-9]{2,12})?", language.lower()):
            return {
                "success": False,
                "summary": None,
                "title": None,
                "url": None,
                "language": language,
                "error": "Invalid language code format"
            }

        import requests

        query_encoded = urllib.parse.quote(query, safe="")
        url = f"https://{language.lower()}.wikipedia.org/api/rest_v1/page/summary/{query_encoded}"
        
        # B3: Wikipedia domain allowlist 檢查
        parsed = urllib.parse.urlparse(url)
        if not (parsed.hostname and parsed.hostname.endswith(WIKIPEDIA_DOMAIN_SUFFIX)):
            return {
                "success": False,
                "summary": None,
                "title": None,
                "url": None,
                "language": language,
                "error": f"Domain not allowed (must be *.wikipedia.org)"
            }
        
        ok, msg = _is_safe_url(url)
        if not ok:
            return {
                "success": False,
                "summary": None,
                "title": None,
                "url": None,
                "language": language,
                "error": f"URL blocked: {msg}"
            }
        
        session = _create_safe_session() or requests.Session()
        response = session.get(url, timeout=10)
        response.raise_for_status()

        data = response.json()

        extract = data.get('extract', '')
        sentences_list = extract.split('. ')[:sentences]
        summary = '. '.join(sentences_list)
        if summary and not summary.endswith('.'):
            summary += '.'

        return {
            "success": True,
            "summary": summary,
            "title": data.get('title', ''),
            "url": data.get('content_urls', {}).get('desktop', {}).get('page', ''),
            "language": language,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "summary": None,
            "title": None,
            "url": None,
            "language": language,
            "error": str(e)
        }


def geocoding(
    location: str,
    return_info: List[str]
) -> Dict[str, Any]:
    """Geocoding with SSRF protection."""
    try:
        import requests

        url = "https://nominatim.openstreetmap.org/search"
        
        ok, msg = _is_safe_url(url, ALLOWED_API_DOMAINS)
        if not ok:
            return {
                "success": False,
                "info": None,
                "location": location,
                "error": f"URL blocked: {msg}"
            }
        
        params = {
            "q": location,
            "format": "json",
            "limit": 1
        }
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36"
        }

        session = _create_safe_session() or requests.Session()
        response = session.get(url, params=params, headers=headers, timeout=10)
        response.raise_for_status()

        data = response.json()

        if not data:
            return {
                "success": False,
                "info": None,
                "location": location,
                "error": "Location not found"
            }

        result = data[0]
        info = {}

        if "coordinates" in return_info:
            info["coordinates"] = {
                "lat": float(result.get("lat", 0)),
                "lon": float(result.get("lon", 0))
            }
        if "country" in return_info:
            info["country"] = result.get("display_name", "").split(", ")[-1]
        if "full_address" in return_info:
            info["full_address"] = result.get("display_name", "")

        return {
            "success": True,
            "info": info,
            "location": location,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "info": None,
            "location": location,
            "error": str(e)
        }


def create_csv(
    data: List[Dict[str, Any]],
    filename: str,
    include_header: bool = True
) -> Dict[str, Any]:
    """Create CSV file."""
    try:
        if not data:
            return {
                "success": False,
                "filename": filename,
                "rows_written": 0,
                "columns": 0,
                "error": "No data to write"
            }

        with open(filename, 'w', newline='', encoding='utf-8') as f:
            fieldnames = list(data[0].keys())
            writer = csv.DictWriter(f, fieldnames=fieldnames)

            if include_header:
                writer.writeheader()

            writer.writerows(data)

        return {
            "success": True,
            "filename": filename,
            "rows_written": len(data),
            "columns": len(fieldnames),
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "filename": filename,
            "rows_written": 0,
            "columns": 0,
            "error": str(e)
        }


def create_markdown(
    title: str,
    sections: List[Dict[str, str]],
    filename: Optional[str] = None
) -> Dict[str, Any]:
    """Create Markdown document."""
    try:
        content = f"# {title}\n\n"

        for section in sections:
            heading = section.get("heading", "")
            text = section.get("content", "")
            level = section.get("level", 2)

            content += f"{'#' * level} {heading}\n\n"
            content += f"{text}\n\n"

        if filename:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)

        return {
            "success": True,
            "content": content,
            "filename": filename,
            "sections_count": len(sections),
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "content": None,
            "filename": filename,
            "sections_count": 0,
            "error": str(e)
        }


def validate_data(
    data: List[Dict[str, Any]],
    rules: Dict[str, Dict[str, Any]]
) -> Dict[str, Any]:
    """Validate data against rules."""
    try:
        errors = []
        valid_count = 0

        for i, item in enumerate(data):
            item_errors = []

            for field, rule in rules.items():
                if field not in item:
                    if rule.get("required", False):
                        item_errors.append(f"Field '{field}' is required")
                    continue

                value = item[field]

                if "type" in rule:
                    expected_type = rule["type"]
                    if expected_type == "integer" and not isinstance(value, int):
                        item_errors.append(f"Field '{field}' must be integer")
                    elif expected_type == "float" and not isinstance(value, (int, float)):
                        item_errors.append(f"Field '{field}' must be number")
                    elif expected_type == "string" and not isinstance(value, str):
                        item_errors.append(f"Field '{field}' must be string")
                    elif expected_type == "email":
                        if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', str(value)):
                            item_errors.append(f"Field '{field}' must be valid email")

                if "min" in rule and isinstance(value, (int, float)):
                    if value < rule["min"]:
                        item_errors.append(f"Field '{field}' must be >= {rule['min']}")
                if "max" in rule and isinstance(value, (int, float)):
                    if value > rule["max"]:
                        item_errors.append(f"Field '{field}' must be <= {rule['max']}")

                if "min_length" in rule and isinstance(value, str):
                    if len(value) < rule["min_length"]:
                        item_errors.append(f"Field '{field}' length must be >= {rule['min_length']}")
                if "max_length" in rule and isinstance(value, str):
                    if len(value) > rule["max_length"]:
                        item_errors.append(f"Field '{field}' length must be <= {rule['max_length']}")

            if item_errors:
                errors.append({
                    "row": i,
                    "errors": item_errors
                })
            else:
                valid_count += 1

        is_valid = len(errors) == 0

        return {
            "success": True,
            "is_valid": is_valid,
            "valid_count": valid_count,
            "invalid_count": len(errors),
            "total_count": len(data),
            "errors": errors[:10],
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "is_valid": False,
            "valid_count": 0,
            "invalid_count": 0,
            "total_count": 0,
            "errors": [],
            "error": str(e)
        }


def compare_data(
    data1: List[Dict[str, Any]],
    data2: List[Dict[str, Any]],
    comparison_type: str = "exact"
) -> Dict[str, Any]:
    """Compare two datasets."""
    try:
        if comparison_type == "exact":
            is_equal = data1 == data2
            differences = []

            if not is_equal:
                for i, (item1, item2) in enumerate(zip(data1, data2)):
                    if item1 != item2:
                        differences.append({
                            "row": i,
                            "data1": item1,
                            "data2": item2
                        })

                if len(data1) != len(data2):
                    differences.append({
                        "type": "length_mismatch",
                        "data1_length": len(data1),
                        "data2_length": len(data2)
                    })

        elif comparison_type == "structural":
            is_equal = True
            differences = []

            if data1 and data2:
                keys1 = set(data1[0].keys()) if data1 else set()
                keys2 = set(data2[0].keys()) if data2 else set()

                if keys1 != keys2:
                    is_equal = False
                    differences.append({
                        "type": "structure_mismatch",
                        "only_in_data1": list(keys1 - keys2),
                        "only_in_data2": list(keys2 - keys1)
                    })

        elif comparison_type == "fuzzy":
            is_equal = len(data1) == len(data2)
            similarity = 0

            if data1 and data2:
                matching = sum(1 for item1, item2 in zip(data1, data2) if item1 == item2)
                similarity = matching / max(len(data1), len(data2))

            differences = [{"similarity": similarity}]
        else:
            differences = []
            is_equal = False

        return {
            "success": True,
            "is_equal": is_equal,
            "comparison_type": comparison_type,
            "differences_count": len(differences),
            "differences": differences[:10],
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "is_equal": False,
            "comparison_type": comparison_type,
            "differences_count": 0,
            "differences": [],
            "error": str(e)
        }


# ============================================================
# 再新增5個工具
# ============================================================

def extract_information(
    text: str,
    extract_type: str,
    keywords: Optional[List[str]] = None,
    pattern: Optional[str] = None
) -> Dict[str, Any]:
    """Extract information from text."""
    try:
        if len(text) > MAX_TEXT_LENGTH:
            return {
                "success": False,
                "extracted": [],
                "count": 0,
                "extract_type": extract_type,
                "error": f"Text too large (max {MAX_TEXT_LENGTH})"
            }

        extracted = []

        if extract_type == "numbers":
            pattern = r'[-+]?(?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d+)?(?:[eE][-+]?\d+)?'
            extracted = re.findall(pattern, text)

        elif extract_type == "dates":
            patterns = [
                r'\d{4}-\d{2}-\d{2}',
                r'\d{1,2}/\d{1,2}/\d{4}',
                r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}'
            ]
            for p in patterns:
                extracted.extend(re.findall(p, text, re.IGNORECASE))

        elif extract_type == "urls":
            pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
            extracted = re.findall(pattern, text)

        elif extract_type == "emails":
            pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            extracted = re.findall(pattern, text)

        elif extract_type == "keywords":
            if not keywords:
                return {
                    "success": False,
                    "extracted": [],
                    "count": 0,
                    "extract_type": extract_type,
                    "error": "Keywords required for keyword extraction"
                }
            for kw in keywords:
                if kw.lower() in text.lower():
                    extracted.append(kw)

        elif extract_type == "sentences":
            extracted = re.split(r'[.!?]+', text)
            extracted = [s.strip() for s in extracted if s.strip()]

        elif extract_type == "custom":
            # B2 修正：使用 regex_search() 統一處理 ReDoS 防護
            if not pattern:
                return {
                    "success": False,
                    "extracted": [],
                    "count": 0,
                    "extract_type": extract_type,
                    "error": "Pattern required for custom extraction"
                }
            
            rs = regex_search(text=text, pattern=pattern, return_all=True)
            if not rs["success"]:
                return {
                    "success": False,
                    "extracted": [],
                    "count": 0,
                    "extract_type": extract_type,
                    "error": rs["error"]
                }
            extracted = rs["matches"]

        return {
            "success": True,
            "extracted": extracted,
            "count": len(extracted),
            "extract_type": extract_type,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "extracted": [],
            "count": 0,
            "extract_type": extract_type,
            "error": str(e)
        }


def count_occurrences(
    data: Union[str, List],
    target: Union[str, List[str]],
    case_sensitive: bool = False,
    count_type: str = "exact"
) -> Dict[str, Any]:
    """Count occurrences with case-insensitive list support."""
    try:
        counts = {}
        targets = [target] if isinstance(target, str) else target

        if isinstance(data, str):
            search_text = data if case_sensitive else data.lower()
            for t in targets:
                search_target = t if case_sensitive else t.lower()
                if count_type == "exact":
                    counts[t] = search_text.count(search_target)
                elif count_type == "word":
                    counts[t] = len(re.findall(r'\b' + re.escape(search_target) + r'\b', search_text))
                elif count_type == "contains":
                    counts[t] = search_text.count(search_target)

        elif isinstance(data, list):
            if case_sensitive:
                for t in targets:
                    counts[t] = data.count(t)
            else:
                lowered = [x.lower() if isinstance(x, str) else x for x in data]
                for t in targets:
                    tl = t.lower() if isinstance(t, str) else t
                    counts[t] = sum(1 for x in lowered if x == tl)

        total = sum(counts.values())

        return {
            "success": True,
            "counts": counts,
            "total": total,
            "count_type": count_type,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "counts": {},
            "total": 0,
            "count_type": count_type,
            "error": str(e)
        }


def compare_values(
    value1: Any,
    value2: Any,
    comparison: str = "equal",
    tolerance: float = 0.0,
    similarity_threshold: float = 0.8
) -> Dict[str, Any]:
    """Compare two values."""
    try:
        result = False
        similarity = None

        if comparison == "equal":
            if isinstance(value1, (int, float)) and isinstance(value2, (int, float)):
                result = abs(value1 - value2) <= tolerance
            else:
                result = value1 == value2

        elif comparison == "not_equal":
            result = value1 != value2

        elif comparison == "greater":
            result = value1 > value2

        elif comparison == "less":
            result = value1 < value2

        elif comparison == "greater_equal":
            result = value1 >= value2

        elif comparison == "less_equal":
            result = value1 <= value2

        elif comparison == "contains":
            if isinstance(value1, (list, str)):
                result = value2 in value1
            else:
                result = False

        elif comparison == "similar":
            if isinstance(value1, str) and isinstance(value2, str):
                similarity = difflib.SequenceMatcher(None, value1, value2).ratio()
                result = similarity >= similarity_threshold

        return {
            "success": True,
            "result": result,
            "comparison": comparison,
            "value1": value1,
            "value2": value2,
            "similarity": similarity,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "result": None,
            "comparison": comparison,
            "value1": value1,
            "value2": value2,
            "similarity": None,
            "error": str(e)
        }


def list_operations(
    list1: List,
    list2: Optional[List] = None,
    operation: str = "unique"
) -> Dict[str, Any]:
    """List operations with deterministic ordering."""
    try:
        op = (operation or "").lower()

        def uniq(seq: List[Any]) -> List[Any]:
            seen = set()
            out = []
            for it in seq:
                k = _keyify(it)
                if k not in seen:
                    seen.add(k)
                    out.append(it)
            return out

        if op == "unique":
            result = uniq(list1)

        elif op in ["intersection", "union", "difference", "symmetric_difference"]:
            if list2 is None:
                return {
                    "success": False,
                    "result": [],
                    "operation": operation,
                    "count": 0,
                    "error": "list2 required for binary operations"
                }

            u1 = uniq(list1)
            u2 = uniq(list2)

            keys2 = set(_keyify(x) for x in u2)
            keys1 = set(_keyify(x) for x in u1)

            if op == "intersection":
                result = [x for x in u1 if _keyify(x) in keys2]

            elif op == "difference":
                result = [x for x in u1 if _keyify(x) not in keys2]

            elif op == "union":
                result = list(u1)
                seen = set(_keyify(x) for x in result)
                for x in u2:
                    k = _keyify(x)
                    if k not in seen:
                        seen.add(k)
                        result.append(x)

            elif op == "symmetric_difference":
                result = [x for x in u1 if _keyify(x) not in keys2]
                seen = set(_keyify(x) for x in result)
                for x in u2:
                    k = _keyify(x)
                    if k not in keys1 and k not in seen:
                        seen.add(k)
                        result.append(x)
        else:
            return {
                "success": False,
                "result": [],
                "operation": operation,
                "count": 0,
                "error": f"Unknown operation: {operation}"
            }

        return {
            "success": True,
            "result": result,
            "operation": operation,
            "count": len(result),
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "result": [],
            "operation": operation,
            "count": 0,
            "error": str(e)
        }


def find_in_text(
    text: str,
    search_terms: Union[str, List[str]],
    context_chars: int = 200,
    max_results: int = 5
) -> Dict[str, Any]:
    """Find in text with early stopping for high-frequency terms."""
    try:
        if len(text) > MAX_TEXT_LENGTH:
            return {
                "success": False,
                "matches": [],
                "count": 0,
                "search_terms": [],
                "error": f"Text too large (max {MAX_TEXT_LENGTH})"
            }

        terms = [search_terms] if isinstance(search_terms, str) else search_terms
        matches = []

        line_starts = _build_line_index(text)

        for term in terms:
            if len(matches) >= max_results:
                break

            pattern = re.compile(re.escape(term))
            for m in pattern.finditer(text):
                pos = m.start()

                start = max(0, pos - context_chars)
                end = min(len(text), pos + len(term) + context_chars)
                context = text[start:end]

                line_number = _find_line_number(pos, line_starts)

                matches.append({
                    "term": term,
                    "position": pos,
                    "context": context,
                    "line_number": line_number,
                    "match_text": term,
                    "span": [pos, pos + len(term)]
                })

                if len(matches) >= max_results:
                    break

        return {
            "success": True,
            "matches": matches,
            "count": len(matches),
            "search_terms": terms,
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "matches": [],
            "count": 0,
            "search_terms": [],
            "error": str(e)
        }


# ============================================================
# CRITICAL: Level 3 必需的2個工具（完全修正版）
# ============================================================

def read_xml(file_path: str, encoding: str = "utf-8") -> Dict[str, Any]:
    """
    Read XML file and convert to dictionary with depth/size limits.
    Essential for gaia_val_l3_006.
    """
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_XML_SIZE:
            return {
                "success": False,
                "data": None,
                "root_tag": None,
                "error": f"XML file too large (>{MAX_XML_SIZE} bytes)"
            }
        
        # 優先使用 defusedxml
        try:
            from defusedxml import ElementTree as ET
        except ImportError:
            import xml.etree.ElementTree as ET
        
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        # C4: 增加深度和節點計數檢查
        node_count = [0]  # 使用 list 作為可變容器
        
        def element_to_dict(element, depth=0):
            """Recursively convert XML element to dict with limits."""
            if depth > MAX_XML_DEPTH:
                raise ValueError(f"XML nesting too deep (max {MAX_XML_DEPTH})")
            
            node_count[0] += 1
            if node_count[0] > MAX_XML_NODES:
                raise ValueError(f"Too many XML nodes (max {MAX_XML_NODES})")
            
            result = {}
            
            if element.attrib:
                result['@attributes'] = element.attrib
            
            if element.text and element.text.strip():
                if len(element) == 0:
                    return element.text.strip()
                result['#text'] = element.text.strip()
            
            children = {}
            for child in element:
                child_data = element_to_dict(child, depth + 1)
                
                if child.tag in children:
                    if not isinstance(children[child.tag], list):
                        children[child.tag] = [children[child.tag]]
                    children[child.tag].append(child_data)
                else:
                    children[child.tag] = child_data
            
            result.update(children)
            
            return result
        
        data = {
            root.tag: element_to_dict(root)
        }
        
        return {
            "success": True,
            "data": data,
            "root_tag": root.tag,
            "error": None
        }
        
    except Exception as e:
        return {
            "success": False,
            "data": None,
            "root_tag": None,
            "error": str(e)
        }


def extract_zip(
    zip_path: str, 
    extract_to: Optional[str] = None,
    password: Optional[str] = None
) -> Dict[str, Any]:
    """
    Extract ZIP file with complete security protection.
    Essential for gaia_val_l3_006.
    """
    extract_dir = extract_to
    
    try:
        zip_size = os.path.getsize(zip_path)
        if zip_size > MAX_ZIP_SIZE:
            return {
                "success": False,
                "extract_path": extract_dir,
                "files": [],
                "count": 0,
                "error": f"ZIP file too large (>{MAX_ZIP_SIZE} bytes)"
            }
        
        if extract_dir is None:
            extract_dir = tempfile.mkdtemp(prefix="gaia_zip_")
        else:
            os.makedirs(extract_dir, exist_ok=True)
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            total_size = sum(info.file_size for info in zip_ref.infolist())
            if total_size > MAX_ZIP_SIZE * 10:
                if extract_to is None and os.path.exists(extract_dir):
                    shutil.rmtree(extract_dir, ignore_errors=True)
                return {
                    "success": False,
                    "extract_path": extract_dir,
                    "files": [],
                    "count": 0,
                    "error": f"Uncompressed size too large (potential zip bomb)"
                }
            
            # 使用安全解壓函數
            _safe_extract_zip(zip_ref, extract_dir, password)
        
        # P0-2: 從實際解壓的目錄掃描文件
        extracted_files = []
        base = Path(extract_dir).resolve()
        
        for file_path in base.rglob("*"):
            if file_path.is_file():
                extracted_files.append({
                    "filename": str(file_path.relative_to(base)).replace("\\", "/"),
                    "path": str(file_path),
                    "size": file_path.stat().st_size,
                    "compressed_size": None
                })
        
        return {
            "success": True,
            "extract_path": extract_dir,
            "files": extracted_files,
            "count": len(extracted_files),
            "error": None
        }
        
    except Exception as e:
        if extract_to is None and extract_dir and os.path.exists(extract_dir):
            shutil.rmtree(extract_dir, ignore_errors=True)
        
        return {
            "success": False,
            "extract_path": extract_dir,
            "files": [],
            "count": 0,
            "error": str(e)
        }


# ============================================================
# Tool Registry - 完整43個工具
# ============================================================

GAIA_TOOLS = {
    "read_pdf": read_pdf,
    "web_search": web_search,
    "read_csv": read_csv,
    "read_image": read_image,
    "calculate": calculate,
    "read_excel": read_excel,
    "image_to_text": image_to_text,
    "web_fetch": web_fetch,
    "filter_data": filter_data,
    "read_json": read_json,
    "read_text_file": read_text_file,
    "read_docx": read_docx,
    "aggregate_data": aggregate_data,
    "sort_data": sort_data,
    "analyze_image": analyze_image,
    "date_calculator": date_calculator,
    "unit_converter": unit_converter,
    "join_data": join_data,
    "deduplicate_data": deduplicate_data,
    "pivot_table": pivot_table,
    "fill_missing": fill_missing,
    "sample_data": sample_data,
    "regex_search": regex_search,
    "string_transform": string_transform,
    "encode_decode": encode_decode,
    "split_join_text": split_join_text,
    "statistical_analysis": statistical_analysis,
    "correlation_analysis": correlation_analysis,
    "moving_average": moving_average,
    "currency_converter": currency_converter,
    "wikipedia_search": wikipedia_search,
    "geocoding": geocoding,
    "create_csv": create_csv,
    "create_markdown": create_markdown,
    "validate_data": validate_data,
    "compare_data": compare_data,
    "extract_information": extract_information,
    "count_occurrences": count_occurrences,
    "compare_values": compare_values,
    "list_operations": list_operations,
    "find_in_text": find_in_text,
    "read_xml": read_xml,
    "extract_zip": extract_zip,
}

GAIA_TOOL_SCHEMAS = {
    "read_pdf": {
        "description": "Read PDF file and extract text content",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to PDF file"},
                "page_numbers": {"type": "array", "items": {"type": "integer"}, "description": "Optional list of page numbers to read"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "content": {"type": ["string", "null"]},
                "pages": {"type": "integer"},
                "metadata": {"type": "object"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "web_search": {
        "description": "Search the web using Serper API or simulation",
        "arguments": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search query"},
                "num_results": {"type": "integer", "description": "Number of results to return"}
            },
            "required": ["query"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "results": {"type": "array"},
                "query": {"type": "string"},
                "count": {"type": "integer"},
                "provider": {"type": "string"},
                "is_simulated": {"type": "boolean"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "read_csv": {
        "description": "Read CSV file and convert to structured data",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to CSV file"},
                "encoding": {"type": "string", "description": "File encoding"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "data": {"type": ["array", "null"]},
                "columns": {"type": "array"},
                "rows": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "read_image": {
        "description": "Read image file and return basic information",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to image file"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "format": {"type": ["string", "null"]},
                "mode": {"type": ["string", "null"]},
                "width": {"type": "integer"},
                "height": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "calculate": {
        "description": "Safely evaluate mathematical expressions",
        "arguments": {
            "type": "object",
            "properties": {
                "expression": {"type": "string", "description": "Mathematical expression to evaluate"}
            },
            "required": ["expression"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "result": {"type": ["number", "null"]},
                "expression": {"type": "string"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "read_excel": {
        "description": "Read Excel file and convert to structured data",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to Excel file"},
                "sheet_name": {"type": ["string", "null"], "description": "Optional sheet name"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "data": {"type": ["array", "null"]},
                "columns": {"type": "array"},
                "rows": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "image_to_text": {
        "description": "Extract text from image using OCR",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to image file"},
                "lang": {"type": "string", "description": "Language code"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "text": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "web_fetch": {
        "description": "Fetch webpage content",
        "arguments": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "URL to fetch"},
                "timeout": {"type": "integer", "description": "Request timeout in seconds"}
            },
            "required": ["url"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "content": {"type": ["string", "null"]},
                "status_code": {"type": ["integer", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "filter_data": {
        "description": "Filter data based on conditions",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "List of dictionaries to filter"},
                "conditions": {"type": "object", "description": "Filter conditions"}
            },
            "required": ["data", "conditions"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "filtered_data": {"type": "array"},
                "count": {"type": "integer"},
                "original_count": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "read_json": {
        "description": "Read JSON file",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to JSON file"},
                "encoding": {"type": "string", "description": "File encoding"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "data": {"type": ["object", "array", "null"]},
                "type": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "read_text_file": {
        "description": "Read plain text file",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to text file"},
                "encoding": {"type": "string", "description": "File encoding"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "content": {"type": ["string", "null"]},
                "lines": {"type": "integer"},
                "characters": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "read_docx": {
        "description": "Read Word document",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to DOCX file"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "content": {"type": ["string", "null"]},
                "paragraphs": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "aggregate_data": {
        "description": "Aggregate data by grouping",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "List of dictionaries to aggregate"},
                "group_by": {"type": "string", "description": "Field to group by"},
                "aggregate_field": {"type": "string", "description": "Field to aggregate"},
                "operation": {"type": "string", "description": "Aggregation operation"}
            },
            "required": ["data", "group_by", "aggregate_field"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "aggregated_data": {"type": ["object", "null"]},
                "operation": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "sort_data": {
        "description": "Sort data by specified field",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "List of dictionaries to sort"},
                "sort_by": {"type": "string", "description": "Field to sort by"},
                "reverse": {"type": "boolean", "description": "Sort in reverse order"}
            },
            "required": ["data", "sort_by"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "sorted_data": {"type": "array"},
                "count": {"type": "integer"},
                "sort_by": {"type": ["string", "null"]},
                "order": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "analyze_image": {
        "description": "Analyze image properties",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to image file"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "format": {"type": ["string", "null"]},
                "mode": {"type": ["string", "null"]},
                "size": {"type": ["array", "null"]},
                "mean_color": {"type": ["array", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "date_calculator": {
        "description": "Calculate dates and date differences",
        "arguments": {
            "type": "object",
            "properties": {
                "start_date": {"type": "string", "description": "Start date"},
                "days_to_add": {"type": "integer", "description": "Number of days to add"},
                "end_date": {"type": ["string", "null"], "description": "End date for difference calculation"},
                "date_format": {"type": "string", "description": "Date format string"}
            },
            "required": ["start_date"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "start_date": {"type": "string"},
                "days_added": {"type": ["integer", "null"]},
                "result_date": {"type": ["string", "null"]},
                "end_date": {"type": ["string", "null"]},
                "days_diff": {"type": ["integer", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "unit_converter": {
        "description": "Convert between units",
        "arguments": {
            "type": "object",
            "properties": {
                "value": {"type": "number", "description": "Value to convert"},
                "from_unit": {"type": "string", "description": "Source unit"},
                "to_unit": {"type": "string", "description": "Target unit"},
                "unit_type": {"type": "string", "description": "Unit category"}
            },
            "required": ["value", "from_unit", "to_unit"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "value": {"type": "number"},
                "from_unit": {"type": "string"},
                "to_unit": {"type": "string"},
                "result": {"type": ["number", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "join_data": {
        "description": "Join two datasets",
        "arguments": {
            "type": "object",
            "properties": {
                "data1": {"type": "array", "description": "First dataset"},
                "data2": {"type": "array", "description": "Second dataset"},
                "key1": {"type": "string", "description": "Join key for first dataset"},
                "key2": {"type": "string", "description": "Join key for second dataset"},
                "join_type": {"type": "string", "description": "Type of join"}
            },
            "required": ["data1", "data2", "key1", "key2"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "data": {"type": "array"},
                "count": {"type": "integer"},
                "join_type": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "deduplicate_data": {
        "description": "Remove duplicate entries from data",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "List to deduplicate"},
                "key": {"type": ["string", "null"], "description": "Optional key field for deduplication"}
            },
            "required": ["data"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "data": {"type": "array"},
                "original_count": {"type": "integer"},
                "final_count": {"type": "integer"},
                "duplicates_removed": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "pivot_table": {
        "description": "Create pivot table from data",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "Input data"},
                "index": {"type": "string", "description": "Row index field"},
                "columns": {"type": "string", "description": "Column field"},
                "values": {"type": "string", "description": "Values field"},
                "aggfunc": {"type": "string", "description": "Aggregation function"}
            },
            "required": ["data", "index", "columns", "values"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "pivot_table": {"type": ["object", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "fill_missing": {
        "description": "Fill missing values in data",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "Input data"},
                "strategy": {"type": "string", "description": "Fill strategy"},
                "value": {"description": "Fill value"}
            },
            "required": ["data", "strategy"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "data": {"type": "array"},
                "filled_count": {"type": "integer"},
                "strategy": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "sample_data": {
        "description": "Sample data randomly",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "Input data"},
                "n": {"type": "integer", "description": "Number of samples"},
                "random_state": {"type": ["integer", "null"], "description": "Random seed"}
            },
            "required": ["data", "n"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "sample": {"type": "array"},
                "sample_size": {"type": "integer"},
                "original_size": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "regex_search": {
        "description": "Search text using regular expressions",
        "arguments": {
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "Text to search"},
                "pattern": {"type": "string", "description": "Regex pattern"},
                "flags": {"type": "string", "description": "Regex flags"}
            },
            "required": ["text", "pattern"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "matches": {"type": "array"},
                "count": {"type": "integer"},
                "pattern": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "string_transform": {
        "description": "Transform string case",
        "arguments": {
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "Text to transform"},
                "operation": {"type": "string", "description": "Transformation operation"}
            },
            "required": ["text", "operation"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "result": {"type": ["string", "null"]},
                "operation": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "encode_decode": {
        "description": "Encode or decode text",
        "arguments": {
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "Text to encode/decode"},
                "operation": {"type": "string", "description": "Operation type"},
                "encoding": {"type": "string", "description": "Encoding type"}
            },
            "required": ["text", "operation"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "result": {"type": ["string", "null"]},
                "operation": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "split_join_text": {
        "description": "Split or join text",
        "arguments": {
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "Text to process"},
                "operation": {"type": "string", "description": "Operation type"},
                "separator": {"type": "string", "description": "Separator"}
            },
            "required": ["text", "operation"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "result": {"type": ["string", "array", "null"]},
                "operation": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "statistical_analysis": {
        "description": "Perform statistical analysis",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "items": {"type": "number"}, "description": "Numeric data"},
                "metrics": {"type": "array", "items": {"type": "string"}, "description": "Metrics to calculate"}
            },
            "required": ["data", "metrics"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "statistics": {"type": ["object", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "correlation_analysis": {
        "description": "Calculate correlation between variables",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "Input data"},
                "x_column": {"type": "string", "description": "X variable column"},
                "y_column": {"type": "string", "description": "Y variable column"},
                "method": {"type": "string", "description": "Correlation method"}
            },
            "required": ["data", "x_column", "y_column"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "correlation": {"type": ["number", "null"]},
                "method": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "moving_average": {
        "description": "Calculate moving average",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "items": {"type": "number"}, "description": "Numeric data"},
                "window_size": {"type": "integer", "description": "Window size"}
            },
            "required": ["data", "window_size"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "moving_average": {"type": ["array", "null"]},
                "window_size": {"type": ["integer", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "currency_converter": {
        "description": "Convert currency amounts",
        "arguments": {
            "type": "object",
            "properties": {
                "amount": {"type": "number", "description": "Amount to convert"},
                "from_currency": {"type": "string", "description": "Source currency code"},
                "to_currency": {"type": "string", "description": "Target currency code"}
            },
            "required": ["amount", "from_currency", "to_currency"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "amount": {"type": "number"},
                "from_currency": {"type": "string"},
                "to_currency": {"type": "string"},
                "result": {"type": ["number", "null"]},
                "rate": {"type": ["number", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "wikipedia_search": {
        "description": "Search Wikipedia and get summary",
        "arguments": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search query"},
                "language": {"type": "string", "description": "Language code"},
                "sentences": {"type": "integer", "description": "Number of sentences"}
            },
            "required": ["query"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "summary": {"type": ["string", "null"]},
                "title": {"type": ["string", "null"]},
                "url": {"type": ["string", "null"]},
                "language": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "geocoding": {
        "description": "Get geographic coordinates for location",
        "arguments": {
            "type": "object",
            "properties": {
                "location": {"type": "string", "description": "Location name"},
                "return_info": {"type": "array", "items": {"type": "string"}, "description": "Information to return"}
            },
            "required": ["location", "return_info"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "location": {"type": ["string", "null"]},
                "latitude": {"type": ["number", "null"]},
                "longitude": {"type": ["number", "null"]},
                "display_name": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "create_csv": {
        "description": "Create CSV file from data",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "Data to write"},
                "filename": {"type": "string", "description": "Output filename"},
                "include_header": {"type": "boolean", "description": "Include header row"}
            },
            "required": ["data", "filename"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "filename": {"type": ["string", "null"]},
                "rows": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "create_markdown": {
        "description": "Create Markdown document",
        "arguments": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Document title"},
                "sections": {"type": "array", "description": "Document sections"},
                "filename": {"type": "string", "description": "Output filename"}
            },
            "required": ["title", "sections"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "filename": {"type": ["string", "null"]},
                "content": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "validate_data": {
        "description": "Validate data against rules",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": "array", "description": "Data to validate"},
                "rules": {"type": "object", "description": "Validation rules"}
            },
            "required": ["data", "rules"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "valid": {"type": "boolean"},
                "errors": {"type": "array"},
                "valid_count": {"type": "integer"},
                "invalid_count": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "compare_data": {
        "description": "Compare two datasets",
        "arguments": {
            "type": "object",
            "properties": {
                "data1": {"type": "array", "description": "First dataset"},
                "data2": {"type": "array", "description": "Second dataset"},
                "comparison_type": {"type": "string", "description": "Type of comparison"}
            },
            "required": ["data1", "data2"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "differences": {"type": "array"},
                "common": {"type": "array"},
                "unique_to_first": {"type": "array"},
                "unique_to_second": {"type": "array"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "extract_information": {
        "description": "Extract structured information from text",
        "arguments": {
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "Text to process"},
                "extract_type": {"type": "string", "description": "Type of extraction"},
                "keywords": {"type": "array", "items": {"type": "string"}, "description": "Keywords to find"},
                "pattern": {"type": "string", "description": "Regex pattern"}
            },
            "required": ["text", "extract_type"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "extracted": {"type": ["array", "object", "null"]},
                "count": {"type": "integer"},
                "extract_type": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "count_occurrences": {
        "description": "Count occurrences of items in data",
        "arguments": {
            "type": "object",
            "properties": {
                "data": {"type": ["string", "array"], "description": "Data to search"},
                "target": {"type": ["string", "array"], "description": "Target to count"},
                "case_sensitive": {"type": "boolean", "description": "Case sensitive search"},
                "count_type": {"type": "string", "description": "Type of counting"}
            },
            "required": ["data", "target"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "count": {"type": ["integer", "object", "null"]},
                "case_sensitive": {"type": ["boolean", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "compare_values": {
        "description": "Compare two values with various comparison methods",
        "arguments": {
            "type": "object",
            "properties": {
                "value1": {"description": "First value"},
                "value2": {"description": "Second value"},
                "comparison": {"type": "string", "description": "Comparison type"},
                "tolerance": {"type": "number", "description": "Numeric tolerance"},
                "similarity_threshold": {"type": "number", "description": "Similarity threshold"}
            },
            "required": ["value1", "value2"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "result": {"type": ["boolean", "number", "null"]},
                "comparison": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "list_operations": {
        "description": "Perform set operations on lists",
        "arguments": {
            "type": "object",
            "properties": {
                "list1": {"type": "array", "description": "First list"},
                "list2": {"type": "array", "description": "Second list"},
                "operation": {"type": "string", "description": "Operation type"}
            },
            "required": ["list1", "operation"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "result": {"type": ["array", "null"]},
                "operation": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "find_in_text": {
        "description": "Find search terms in text with context",
        "arguments": {
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "Text to search"},
                "search_terms": {"type": ["string", "array"], "description": "Terms to find"},
                "context_chars": {"type": "integer", "description": "Context characters"},
                "max_results": {"type": "integer", "description": "Maximum results"}
            },
            "required": ["text", "search_terms"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "results": {"type": "array"},
                "count": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "read_xml": {
        "description": "Read and parse XML file",
        "arguments": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to XML file"},
                "encoding": {"type": "string", "description": "File encoding"}
            },
            "required": ["file_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "data": {"type": ["object", "null"]},
                "root_tag": {"type": ["string", "null"]},
                "error": {"type": ["string", "null"]}
            }
        }
    },
    "extract_zip": {
        "description": "Extract ZIP archive",
        "arguments": {
            "type": "object",
            "properties": {
                "zip_path": {"type": "string", "description": "Path to ZIP file"},
                "extract_to": {"type": "string", "description": "Extraction destination"},
                "password": {"type": "string", "description": "ZIP password"}
            },
            "required": ["zip_path"]
        },
        "results": {
            "type": "object",
            "properties": {
                "success": {"type": "boolean"},
                "extract_path": {"type": ["string", "null"]},
                "files": {"type": "array"},
                "count": {"type": "integer"},
                "error": {"type": ["string", "null"]}
            }
        }
    }
}


if __name__ == "__main__":
    print("="*70)
    print("GAIA Tools v2.3.3 Production")
    print("="*70)
    print(f"\n✅ Total tools: {len(GAIA_TOOLS)}")
    print(f"\n🔒 Security Protections (with documented limitations):")
    print(f"  ✅ B1: Symlink blocking (before is_dir check)")
    print(f"  ✅ B2: ReDoS protection (extract_information uses regex_search)")
    print(f"  ✅ C1: Resource cleanup (web_fetch with context manager)")
    print(f"  ✅ C2: Extended IP blocklist (CGNAT, TEST-NET, multicast, etc.)")
    print(f"  ✅ C4: XML depth/node limits (max {MAX_XML_DEPTH} depth, {MAX_XML_NODES} nodes)")
    print(f"  ✅ D1: CSV/Excel size limits ({MAX_CSV_EXCEL_SIZE} bytes)")
    print(f"  ✅ D2: Path checking with Path.relative_to()")
    print(f"  ✅ D3: Content-Type validation (empty = unsupported)")
    print(f"\n📋 Known Limitations:")
    print(f"  ⚠️  SSRF: DNS rebinding / TOCTOU not fully mitigated (best-effort)")
    print(f"  ⚠️  ReDoS: Pattern/text limits reduce risk; no timeout enforcement")
    print(f"  ⚠️  Wikipedia: Domain allowlist (*.wikipedia.org)")
    print(f"\n🔧 SSRF Policy: {'Strict (fail-closed)' if STRICT_SSRF else 'Relaxed (fail-open)'}")
    print(f"\n✨ GAIA Level 3: All 10 tasks fully supported")
